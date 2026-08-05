import os
import shutil
import zipfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
RELEASE_DIR_NAME = "HCM_HUB_Dashboard_v2_5_Release_Stable"
RELEASE_DIR = os.path.join(ROOT_DIR, RELEASE_DIR_NAME)
ZIP_FILE = os.path.join(ROOT_DIR, f"{RELEASE_DIR_NAME}.zip")

# Clean target release directory
if os.path.exists(RELEASE_DIR):
    shutil.rmtree(RELEASE_DIR)
os.makedirs(RELEASE_DIR, exist_ok=True)

print(f"[RELEASE BUNDLE] Creating Release Bundle Folder: {RELEASE_DIR}")

# 1. Copy Backend ETL & Configs
backend_dest = os.path.join(RELEASE_DIR, "backend_sync")
os.makedirs(backend_dest, exist_ok=True)
for item in ["sync_postgre.py", "pipeline_unified_v6.py", "setup_sync_postgre_task.ps1", "configs"]:
    src = os.path.join(ROOT_DIR, "backend_sync", item)
    dst = os.path.join(backend_dest, item)
    if os.path.isdir(src):
        shutil.copytree(src, dst, dirs_exist_ok=True)
    elif os.path.isfile(src):
        shutil.copy2(src, dst)

# 2. Copy Frontend Source
src_dest = os.path.join(RELEASE_DIR, "src")
os.makedirs(src_dest, exist_ok=True)
if os.path.exists(os.path.join(ROOT_DIR, "src")):
    shutil.copytree(os.path.join(ROOT_DIR, "src"), src_dest, dirs_exist_ok=True)

# 3. Copy Data Snapshots (data and public/data)
data_dest = os.path.join(RELEASE_DIR, "data")
public_dest = os.path.join(RELEASE_DIR, "public", "data")
os.makedirs(data_dest, exist_ok=True)
os.makedirs(public_dest, exist_ok=True)

if os.path.exists(os.path.join(ROOT_DIR, "data")):
    shutil.copytree(os.path.join(ROOT_DIR, "data"), data_dest, dirs_exist_ok=True)
if os.path.exists(os.path.join(ROOT_DIR, "public", "data")):
    shutil.copytree(os.path.join(ROOT_DIR, "public", "data"), public_dest, dirs_exist_ok=True)

# 4. Copy Root Files
for rf in ["package.json", "vite.config.ts", "index.html", "tsconfig.json"]:
    src_rf = os.path.join(ROOT_DIR, rf)
    if os.path.exists(src_rf):
        shutil.copy2(src_rf, os.path.join(RELEASE_DIR, rf))

# 5. Generate Markdown Technical Document v2.5
md_path = os.path.join(RELEASE_DIR, "Tai_Lieu_Ky_Thuat_Dashboard_HCM_HUB_v2_5.md")
md_content = """# TÀI LIỆU KỸ THUẬT VẬN HÀNH & KIẾN TRÚC HỆ THỐNG
## DASHBOARD BÁO CÁO KHO HCM HUB (J&T CARGO) - VERSION 2.5 STABLE
**Quy trình End-to-End: JFS API → PostgreSQL (28 Cột) → JSON Realtime → React UI**

---

### 1. TỔNG QUAN KIẾN TRÚC & NGUYÊN TẮC VẬN HÀNH
Hệ thống Dashboard HCM HUB hoạt động theo kiến trúc ETL phân tán kết hợp CSDL PostgreSQL Local và hosting tĩnh GitHub Pages:
- **Phase 1 (Data Ingestion):** `pipeline_unified_v6.py` thực hiện kéo dữ liệu song song 7 ngày từ 8 nguồn JFS API. Mức độ song song được tối ưu với `PAGE_WORKERS = 20` (kéo 20 trang API đồng thời) và `POOL_SIZE = 40`.
- **Phase 2 (PostgreSQL Data Processing & Micro-JSON Exporter):** `sync_postgre.py` lưu trữ 28 cột chuẩn hóa vào `enriched.dispatch_enriched` trong PostgreSQL (cổng 5433). Đồng thời thực hiện xuất đồng bộ các gói dữ liệu Micro-JSON (`inbound_kpi_summary.json`, `inbound_truck_eta.json`, `inventory.json`, `last_update.json`...) vào **4 vị trí đồng thời** (`data/`, `data/live/`, `public/data/`, `public/data/live/`) để phục vụ cả Localhost lẫn GitHub Pages Build.
- **Phase 3 (Git Synchronization & Retention):** Tự động đẩy Git commit với danh sách `ROLLING_FILES` chứa 35+ file dữ liệu rolling và thực hiện dọn dẹp dữ liệu PostgreSQL retention 90 ngày (`DELETE FROM raw.scan_logs WHERE scan_time < CURRENT_DATE - 90`).

---

### 2. QUY TẮC NGHIỆP VỤ ĐÃ CHỐT VÀ NÂNG CẤP TRONG VERSION 2.5

#### 2.1 Màn Hình Operational Monitor (Master Layout)
- **Top Card Title & Tỉ Lệ Lấp Đầy (Capacity Fill Rate %):**
  - Tiêu đề thẻ trên cùng được chốt chuẩn là **"TỈ LỆ LẤP ĐẦY"**.
  - Công thức tính:
    $$\text{Tỉ Lệ Lấp Đầy} = \left(\frac{\sum \text{tCur}}{\sum \text{tCap}}\right) \times 100$$
    Trong đó $\sum \text{tCur}$ là tổng số đơn hiện có trên các kệ và $\sum \text{tCap}$ là tổng sức chứa của toàn bộ các kệ bãi.
- **Thẻ Tỉ Lệ Outbound:**
  - Tiêu đề thẻ phía dưới là **"Tỉ lệ Outbound"**.
  - Công thức tính:
    $$\text{outboundRate} = \frac{\text{tCur}}{\text{tCur} + \text{tBacklog}}$$
- **Thanh Chữ Chạy Realtime (Ticker Bar):**
  - Tích hợp hiệu ứng cuộn mượt mà `@keyframes inlineMarquee` trực tiếp trong CSS.
  - Tự động dịch chuyển `left` offset theo trạng thái Sidebar (`left-12` khi thu gọn 48px, `left-40` khi mở rộng 160px) để không bao giờ bị đè hoặc khuất chữ.

#### 2.2 Màn Hình Inbound Dashboard
- **Truck Forecast ETA +36h Transport Rule:**
  - Tất cả các xe Linehaul xuất phát từ **BN HUB** đều được áp dụng quy tắc cộng **+36 tiếng** vào mốc giờ xuất bãi (`ETA = actualDeparture + 36h`).
  - Giữ lại danh sách các xe đã quét cổng HUB (`actualArrival` / `shipmentState = 4`) nhưng chưa quét inbound hàng hóa (`flag_inbound = 0`) trong danh sách xe chờ nhập kho ngày Hôm nay (05/08) và Ngày mai (06/08).
- **KPI Volume Card Filter:**
  - Trạng thái hợp lệ cho thẻ Volume đạt chuẩn exact ~13.1k đơn bằng cách bao gồm đủ 4 trạng thái: `'Inbound'`, `'Transporting'`, `'Pickup Done'`, `'Created'`.

---

### 3. HƯỚNG DẪN KÍCH HOẠT & THEO DÕI

1. **Khởi chạy Luồng Đồng Bộ Tự Động (Kích hoạt Task Scheduler 30 phút/lần):**
   ```powershell
   powershell -ExecutionPolicy Bypass -File ./backend_sync/setup_sync_postgre_task.ps1
   ```
2. **Khởi chạy Localhost Dev Server:**
   ```bash
   npx vite --host 0.0.0.0 --port 5173
   ```
3. **Chạy Luồng Đồng Bộ Bằng Tay Ngay Lập Tức:**
   ```bash
   python backend_sync/sync_postgre.py
   ```
"""

with open(md_path, "w", encoding="utf-8") as f:
    f.write(md_content)

print(f"[RELEASE BUNDLE] Generated Markdown Technical Document: {md_path}")

# 6. Generate Runbook Manual
runbook_path = os.path.join(RELEASE_DIR, "HUONG_DAN_VAN_HANH_SYSTEM_RUNBOOK.md")
runbook_content = """# HƯỚNG DẪN VẬN HÀNH & KỊCH BẢN BẢO TRÌ (RUNBOOK)
## HỆ THỐNG DASHBOARD HCM HUB J&T CARGO - RELEASE v2.5 STABLE

---

### 1. QUY TRÌNH CHẠY TỰ ĐỘNG (BACKGROUND SYNC)
- **Tần suất:** Mỗi 30 phút hệ thống Windows Task Scheduler (`Sync_Postgre_30m`) sẽ tự động kích hoạt script `backend_sync/sync_postgre.py`.
- **Tốc độ xử lý:** Luồng kéo 7 ngày JFS API song song với 20 luồng kéo trang (`PAGE_WORKERS = 20`) hoàn tất chỉ trong 1.5 - 3 phút.
- **Log theo dõi:** Mọi nhật ký chạy được ghi lại tại `backend_sync/logs/` hoặc kiểm tra thẻ `Update: HH:MM:SS` ở góc phải Header Dashboard.

---

### 2. KIỂM TRA & XỬ LÝ SỰ CỐ NHANH (TROUBLESHOOTING)

| Sự Cố | Nguyên Nhân | Cách Xử Lý Nhanh |
| :--- | :--- | :--- |
| **Báo cáo trên GitHub Pages chưa đổi** | Trình duyệt lưu Cache HTTP | Bấm `Ctrl + Shift + R` (Hard Refresh) hoặc mở cửa sổ Ẩn danh. |
| **Token JFS API hết hạn** | Tài khoản JFS bị logout | Thao tác đăng nhập lại trên JFS web hoặc xóa cache token trong script. |
| **Vite Dev Server lỗi JS Syntax** | Cache `.vite` bị kẹt | Chạy `npx vite --force --port 5173` để làm sạch cache. |
| **PostgreSQL không kết nối được** | Dịch vụ Postgres 5433 ngắt | Mở Services Windows và Restart `postgresql-x64-16`. |

---

### 3. ĐƯỜNG DẪN TRUY CẬP HỆ THỐNG
- Localhost: `http://localhost:5173/`
- GitHub Pages: `https://lehoangtienpham2395.github.io/sortation-center-layout/`
- Repository GitHub: `https://github.com/lehoangtienpham2395/sortation-center-layout`
"""

with open(runbook_path, "w", encoding="utf-8") as f:
    f.write(runbook_content)

print(f"[RELEASE BUNDLE] Generated Runbook Manual: {runbook_path}")

# 7. Generate Word Document (.docx) Technical Spec v2.5
docx_path = os.path.join(RELEASE_DIR, "Tai_Lieu_Ky_Thuat_Dashboard_HCM_HUB_v2_5.docx")

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

C_TITLE = RGBColor(0x1F, 0x38, 0x64)
C_H1 = RGBColor(0x1F, 0x38, 0x64)
C_H2 = RGBColor(0x2E, 0x74, 0xB5)
C_BODY = RGBColor(0x00, 0x00, 0x00)
C_GREY = RGBColor(0x59, 0x59, 0x59)
FONT = "Times New Roman"

def add_t(text, size=20, bold=True, color=C_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold; r.font.name = FONT; r.font.size = Pt(size); r.font.color.rgb = color
    return p

def add_h1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.clear()
    r = p.add_run(text)
    r.bold = True; r.font.name = FONT; r.font.size = Pt(14); r.font.color.rgb = C_H1
    return p

def add_h2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.clear()
    r = p.add_run(text)
    r.bold = True; r.font.name = FONT; r.font.size = Pt(12); r.font.color.rgb = C_H2
    return p

def add_b(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.font.name = FONT; r.font.size = Pt(11); r.font.color.rgb = C_BODY
    return p

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(11)
    return p

def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_table(headers, rows_data, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            t.columns[i].width = Cm(w)
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        shade_cell(cell, "1F3864")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True; r.font.name = FONT; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, r_data in enumerate(rows_data):
        row = t.add_row()
        fill = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_text in enumerate(r_data):
            cell = row.cells[c_idx]
            shade_cell(cell, fill)
            p = cell.paragraphs[0]
            r = p.add_run(cell_text)
            r.font.name = FONT; r.font.size = Pt(10)
    return t

add_t("TÀI LIỆU KỸ THUẬT & KIẾN TRÚC HỆ THỐNG", size=18, color=C_TITLE)
add_t("DASHBOARD BÁO CÁO KHO HCM HUB (J&T CARGO)", size=15, color=C_H2)
add_t("Phiên bản v2.5 Stable  |  Chốt Ngày: 05/08/2026", size=11, bold=False, color=C_GREY)
doc.add_paragraph()

add_h1("1. TỔNG QUAN HỆ THỐNG & ĐỒNG BỘ 4 VỊ TRÍ DATA")
add_b("Hệ thống Dashboard HCM HUB v2.5 hoạt động theo kiến trúc ETL phân tán song song, tự động kéo dữ liệu JFS API mỗi 30 phút, xử lý lưu trữ vào PostgreSQL và đồng bộ dữ liệu vào 4 thư mục đích:")
add_bullet("1. data/ (Dữ liệu Live Localhost)")
add_bullet("2. data/live/ (Snapshot Live Micro-JSON)")
add_bullet("3. public/data/ (Dữ liệu phục vụ Build GitHub Pages)")
add_bullet("4. public/data/live/ (Micro-JSON cho GitHub Pages Build)")

add_h1("2. QUY TẮC THIẾT KẾ UI & CÔNG THỨC CHUẨN")
add_table(
    ["Màn Hình / Card", "Quy Tắc & Công Thức", "Trạng Thái"],
    [
        ["Operational Monitor Top Card", "TỈ LỆ LẤP ĐẦY = (Sum(tCur) / Sum(tCap)) * 100", "Đã chốt chuẩn"],
        ["Operational Monitor Lower Card", "Tỉ lệ Outbound = tCur / (tCur + tBacklog)", "Đã chốt chuẩn"],
        ["Inbound Truck ETA", "BN HUB Linehaul Truck +36h Transport Time Offset", "Đã chốt chuẩn"],
        ["Header Marquee Ticker Bar", "@keyframes inlineMarquee, offset left-12 vs left-40", "Đã chốt chuẩn"],
    ],
    col_widths=[6, 9, 3]
)

doc.save(docx_path)
print(f"[RELEASE BUNDLE] Generated Word Technical Document: {docx_path}")

# 8. Create Zip Archive
print(f"[RELEASE BUNDLE] Zipping release bundle into {ZIP_FILE}...")
with zipfile.ZipFile(ZIP_FILE, 'w', zipfile.ZIP_DEFLATED) as zipf:
    for root, dirs, files in os.walk(RELEASE_DIR):
        for file in files:
            full_path = os.path.join(root, file)
            rel_path = os.path.relpath(full_path, ROOT_DIR)
            zipf.write(full_path, rel_path)

print(f"[RELEASE BUNDLE] RELEASE BUNDLE CREATED SUCCESSFULLY: {ZIP_FILE}")
