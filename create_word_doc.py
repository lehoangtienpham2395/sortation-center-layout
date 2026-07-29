import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = create_element('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = create_element('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = create_element(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(16, 44, 87) # Deep Navy
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 90, 150) # Steel Blue
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(50, 50, 50)
    return p

def add_body_p(doc, text="", bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
    return p

def add_bullet_p(doc, title, text=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(title)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.bold = True
    if text:
        r2 = p.add_run(" " + text)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(11)
    return p

def add_callout(doc, text, title="GHI CHÚ KỸ THUẬT:"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F0F4F8")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = create_element('w:tcBorders')
    left = create_element('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24') # 3pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), '1E5A96')
    tcBorders.append(left)
    for b in ['top', 'right', 'bottom']:
        node = create_element(f'w:{b}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r_title = p.add_run(title + " ")
    r_title.bold = True
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(30, 90, 150)

    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(10.5)
    r_text.font.color.rgb = RGBColor(40, 40, 40)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def format_table_headers_and_cells(table, col_widths, headers, data):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "1E5A96")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Segoe UI'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Data rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.add_row().cells
        bg_color = "F9FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = str(cell_value)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=120, right=120)
            p = row_cells[c_idx].paragraphs[0]
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(40, 40, 40)

    # Set column widths
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)

def generate_docx():
    doc = Document()

    # Set Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = title_p.add_run("TÀI LIỆU KỸ THUẬT VẬN HÀNH & KIẾN TRÚC HỆ THỐNG\nDASHBOARD BÁO CÁO KHO HCM HUB (JT CARGO)")
    r_title.font.name = 'Segoe UI'
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(16, 44, 87)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(16)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run("Quy trình End-to-End: JFS API → PostgreSQL (28 Cột) → JSON Payload → React Frontend")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 1
    add_heading_1(doc, "1. TỔNG QUAN KIẾN TRÚC HỆ THỐNG (END-TO-END)")
    add_body_p(doc, "Hệ thống Dashboard báo cáo vận hành kho HCM HUB hoạt động theo mô hình ETL phân tán song song, kết nối dữ liệu trực tiếp từ JFS API (gw.jtcargo.com.vn) qua cơ sở dữ liệu PostgreSQL local, xuất các gói JSON nén và đồng bộ lên GitHub Pages hiển thị cho người dùng cuối.")
    
    add_callout(doc, 
        "JFS API (7 Nguồn Song Song)\n"
        "   ↓ (pipeline_unified_v6.py — Python ETL Engine)\n"
        "PostgreSQL (Database logistics_db — Schema enriched.dispatch_enriched - 28 Cột)\n"
        "   ↓ (sync_postgre.py — JSON Exporter & Data Aggregator | Windows Task Scheduler 'Sync_Postgre_30m' 30m/lần)\n"
        "Thư mục data/*.json & latest.json.gz (Repository Data Files)\n"
        "   ↓ (git push main + GitHub Actions Trigger theo paths)\n"
        "GitHub Pages CDN → React UI App (App.tsx / InboundDashboard / HeatmapDashboard)",
        "SƠ ĐỒ LUỒNG DỮ LIỆU CHÍNH:"
    )

    # Section 2
    add_heading_1(doc, "2. CÁC NGUYÊN TẮC KỸ THUẬT CỐT LÕI (ĐÃ CHUẨN HÓA)")
    principles = [
        ("1. Chu kỳ Vận hành (06:00 - 06:00):", "Hàm `get_op_date()` trong Python quy đổi mọi mốc thời gian (created_time, inbound_scandate, arrival_scandate, pickup_time) về đúng ca làm việc. Mọi mốc giờ từ 00:00:00 - 05:59:59 sẽ tự động tính lùi 1 ngày để thuộc ca đêm hôm trước."),
        ("2. Khóa chính & Khử trùng lặp (Primary Key & Dedup):", "Python dùng `seen_wb` set khử trùng ngay khi đọc Dispatch. PostgreSQL thiết lập cột `tracking` làm PRIMARY KEY. Đảm bảo mỗi đơn hàng chỉ tồn tại 1 dòng duy nhất trong DB, tránh nhân đôi sản lượng."),
        ("3. Cập nhật ghi đè (Upsert Logic):", "Dùng câu lệnh `INSERT INTO ... ON CONFLICT (tracking) DO UPDATE SET` cho phép tự động cập nhật trạng thái mới nhất của vận đơn khi kiện hàng trôi qua các nấc vận hành mà không tạo dòng rác."),
        ("4. Mốc thời gian quét mới nhất (Latest Timestamp Mapping):", "Sử dụng `max(scanDate)` khi ghép nối các bảng Inbound, Outbound và Arrival Scan để bảo đảm luôn lấy mốc quét mới nhất nếu kiện hàng bị bắn mã nhiều lần."),
        ("5. Thứ tự Ưu tiên Trạng thái Tồn kho (inv_status):", "Quy định thứ tự ưu tiên: `Outbound` (Đã xuất) > `Inbound` (Đang tại bãi) > `Transporting` (Đang trên đường) > `Created` (Mới tạo). Đơn đã Outbound không bao giờ bị tính trùng ở khâu nhập kho."),
        ("6. Phân loại cờ Vận hành & Baseline Rớt Hôm Trước 06:00 AM:", "Chốt cố định Baseline `rot_hom_truoc` lúc 06:00 AM trong bảng `enriched.daily_baseline_snapshot` cho Inbound Dashboard (giữ nguyên không giảm để làm báo cáo ca). Đồng thời duy trì `rot_hom_truoc_live` giảm động theo thời gian thực trên Thẻ Volume Layout Dashboard cho team Outbound xử lý."),
        ("7. Chuẩn hóa Mã Bưu cục (Master Data Driven):", "Đọc `valid.csv` để tra cứu `dict_zone`, `dict_area`, `dict_station`. Ưu tiên lookup bằng `dispatch_code` (sortcode 10 ký tự), nếu rỗng fallback sang `next_station`. Toàn bộ cấu trúc khu vực, sức chứa (capacity) được đọc tự động 100% từ Master Config, loại bỏ hoàn toàn việc ghi đè cứng trong code."),
        ("8. Trích xuất Trạm nguồn & Region Tagging (is_north/region):", "Trích xuất trực tiếp `upOrNextStation` và `sendSite`. Backend tự động gắn thuộc tính `is_north` và `region` ('north'/'south') cho từng record trong `inbound.json` từ Master Config (`valid.csv` -> `dict_zone`). Đã loại bỏ hoàn toàn danh sách cứng `NORTH_POST_OFFICES` và 11 lần gọi `isNorthStation` ở Frontend."),
        ("9. Cấu trúc Gom cụm Aggregate (15-Tuple Key):", "Gom nhóm dữ liệu theo 15-tuple key (bưu cục, trạng thái, 4 ngày vận hành, 4 mốc giờ, drop_type, trip_code, thời gian xe, is_rebound). Cộng dồn `volume += 1` và `weight_kg += orders_weight` để nén file JSON siêu nhẹ mà vẫn giữ đủ độ phân giải phân tích."),
        ("10. Quy chuẩn Đơn vị Trọng lượng Đồng nhất (Tấn - Single Source of Truth):", "Chỉ số trọng lượng được quy chuẩn lưu trữ duy nhất dưới dạng Tấn từ Backend ETL (`weight_ton = weight_kg / 1000`). Frontend đọc thẳng số từ JSON và append chuỗi 'Tấn' lên UI, loại bỏ hoàn toàn các hàm quy đổi 2 lần hay đoán đơn vị rủi ro."),
        ("11. Phân slot Giờ Heatmap (Hourly Bucket):", "Tạo 24 slot giờ (`00:00` - `23:00`), lọc các đơn đã Inbound có `op_date_inbound == today` để đếm tần suất nhập kho theo khung giờ cao điểm."),
        ("12. Chuẩn hóa Trạng thái từ Nguồn Backend (clean_status_sys):", "Triển khai hàm `clean_status_sys()` trong `sync_postgre.py` quy đổi tất cả alias thô từ API về bộ 5 enum chuẩn (`Inbound`, `Transporting`, `Pickup Done`, `Created`, `Outbound`) ngay tại tầng Backend ETL trước khi đẩy vào PostgreSQL."),
        ("13. Chế độ Môi trường Linh hoạt (Dual Hostname Fetching):", "Tự động phát hiện môi trường: Nếu chạy trên `github.io` sẽ fetch CDN Raw GitHub Pages kèm parameter timestamp (`?t=Date.now()`); nếu chạy local sẽ fetch `./data/`."),
        ("14. Tải nén File latest.json.gz qua Native DecompressionStream & Smart Polling:", "Backend nén `inbound.json` (37MB) thành `latest.json.gz` (271KB). Frontend React sử dụng API native `DecompressionStream('gzip')` giải nén trực tiếp trên trình duyệt, giảm 98% dung lượng tải lần đầu. Đồng thời triển khai Smart Polling 60s kiểm tra `last_update.json` trước khi refetch để bảo vệ CDN cache."),
        ("15. Tự động hóa Tiến trình ETL bằng Windows Task Scheduler (30 phút/lần):", "Cấu hình Task Scheduler tự động kích hoạt `sync_postgre.py` định kỳ 30 phút/lần qua task `Sync_Postgre_30m`. Sử dụng wrapper script `run_sync_postgre.bat` thiết lập môi trường UTF-8, chuyển working directory chuẩn và tự động ghi log hoạt động chi tiết vào `sync_postgre.log`. Đồng thời tích hợp kịch bản khởi tạo PowerShell `setup_sync_postgre_task.ps1` để tự động hóa toàn bộ quy trình triển khai dịch vụ 24/7."),
        ("16. Hợp đồng Dữ liệu Trung tâm (Data Contract Single Source of Truth):", "Triển khai đặc tả hợp đồng dữ liệu tập trung `data_contract.json` (Backend) và `data_contract.ts` (Frontend) đóng vai trò là Nguồn Chân Lý Duy Nhất. Định nghĩa cứng tập enum trạng thái Display-Ready (`Inbound`, `Transporting`, `Pickup Done`, `Created`, `Outbound`), quy chuẩn đơn vị `weight_ton` (Tấn, hiển thị trực tiếp không tính toán lại ở Frontend) và bảng ánh xạ key. Đặc biệt: Phân định rõ ràng `station_name` ↔ `Bưu cục đích` (định hướng ô chứa/tuyến) với `pickup_station` ↔ `Bưu cục nộp/lấy hàng` (trạm nguồn) tránh nhầm lẫn giữa 2 đầu bưu cục. Backend chạy kiểm tra `validate_payload_contract()` trước khi xuất file JSON, ngăn ngừa triệt để mọi rủi ro tái phát lỗi sai format hay lệch enum."),
        ("17. Tối ưu hóa Gom cụm Pre-Aggregation & Giảm 83% Dung lượng Payload:", "Tối ưu hóa thuật toán gom nhóm 15-tuple key trong `sync_postgre.py`: Với các đơn đã `Inbound` thành công, thu gọn các mốc giờ lịch sử cũ; với các đơn chưa `Inbound`, bucket thời gian theo khung GIỜ `HH:00:00` thay vì định dạng PHÚT (`:16`). Kết quả làm giảm số dòng trong `inbound.json` từ 13,264 dòng xuống chỉ còn 2,493 dòng (giảm 81% số dòng), giảm dung lượng file thô từ 8.3MB xuống 1.4MB (giảm 83%), và dung lượng file nén `latest.json.gz` giảm từ 271KB xuống chỉ còn 34KB siêu nhẹ, giúp trang web load cực nhanh trên mọi thiết bị.")
    ]

    for title, text in principles:
        add_bullet_p(doc, title, text)

    # Section 3
    add_heading_1(doc, "3. CHI TIẾT CÁC CỘT TẢI TỪ JFS API & LOGIC XỬ LÝ PYTHON")
    add_body_p(doc, "Bảng mổ xẻ dữ liệu chi tiết từ khi kéo từ 7 API JFS về cho đến khi được xử lý bằng Python:")

    jfs_headers = ["Tên Cột Thô JFS API", "Nguồn API", "Cách Lọc & Xử Lý Python", "Tác Dụng Nghiệp Vụ"]
    jfs_widths = [1.8, 1.3, 2.3, 1.8]
    jfs_data = [
        ["waybillId / waybillNo", "Dispatch", "Dùng clean_wb() xóa đuôi .0, bỏ space, đổi HOA. Đưa vào seen_wb set để dedup.", "Mã vận đơn (Primary Key)"],
        ["inputTime / dispatchNetworkTime", "Dispatch", "Lấy YYYY-MM-DD HH:MM:SS. Bỏ qua nếu rỗng.", "Thời gian tạo đơn / phân phối"],
        ["terminalDispatchCode", "Dispatch", "Regex [A-Z]{2,3}\\d{3}[A-Z0-9] rút gọn thành mã 10 ký tự (VD: HCM004H).", "Mã phân loại (Sortcode)"],
        ["orderStatusName", "Dispatch", "Lọc bỏ toàn bộ đơn có orderStatusName == 'Da huy'.", "Trạng thái đơn OMS gốc"],
        ["packageChargeWeight", "Dispatch", "float(val or 0.0). Tính bằng Gram.", "Trọng lượng tính cước gốc"],
        ["packageNumber", "Dispatch", "int(val or 1). Mặc định 1 kiện.", "Số lượng kiện hàng"],
        ["pickNetworkName", "Dispatch", ".strip(). Trích xuất tên bưu cục lấy hàng.", "Bưu cục nộp/lấy hàng gốc"],
        ["realPickNetworkName", "Dispatch", ".strip(). Tên bưu cục lấy hàng thực tế.", "Bưu cục lấy thực tế"],
        ["pickTime", "Dispatch", "Mốc thời gian lấy hàng thành công.", "Thời gian lấy hàng (Pickup)"],
        ["proxyAreaCode / flowTypeDesc", "Dispatch", "Trích xuất mã vùng và mô tả luồng hàng.", "Mã khu vực & Luồng hàng"],
        ["billNo / scanDate / transferCode", "Inbound Scan", "Gom nhóm theo billNo, dùng max(scanDate) chỉ lấy mốc nhập kho mới nhất.", "Mốc quét Nhập kho HUB & Trip Code"],
        ["billNo / scanDate", "Outbound Scan", "Gom nhóm theo billNo, dùng max(scanDate) chỉ lấy mốc xuất kho mới nhất.", "Mốc quét Xuất kho HUB"],
        ["billcode / scantime / transfercode", "Arrival Scan", "Lọc theo bưu cục gửi về HCM004H. Dùng max(scantime) lấy mốc mới nhất.", "Mốc đơn đến bưu cục gửi"],
        ["actualDepartureTime / actualArrivalTime", "Linehaul/Shuttle", "Ánh xạ qua trip_code tạo dict ttm[trip_code] lấy mốc giờ xe xuất/đến bến.", "Thời gian xe chạy / xe đến"]
    ]
    tbl_jfs = doc.add_table(rows=1, cols=4)
    format_table_headers_and_cells(tbl_jfs, jfs_widths, jfs_headers, jfs_data)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 4
    add_heading_1(doc, "4. CHI TIẾT 28 CỘT TRONG POSTGRESQL (enriched.dispatch_enriched)")
    add_body_p(doc, "Bảng cơ sở dữ liệu PostgreSQL lưu trữ toàn bộ dữ liệu 7 ngày gần nhất với 28 cột được thiết kế tối ưu:")

    pg_headers = ["STT", "Tên Cột PostgreSQL", "Kiểu Dữ Liệu", "Logic Sinh Cột Python / SQL", "Tác Dụng Nghiệp Vụ"]
    pg_widths = [0.5, 1.8, 1.1, 2.3, 1.5]
    pg_data = [
        ["1", "tracking", "TEXT (PK)", "clean_wb(waybillId)", "Khóa chính (Primary Key), bảo đảm duy nhất"],
        ["2", "data_source", "TEXT", "Hằng số 'pipeline_v6'", "Ghi nhận nguồn ETL"],
        ["3", "status_sys", "TEXT", "clean_status_sys(orderStatusName)", "Trạng thái chuẩn hoá (Inbound, Pickup Done...)"],
        ["4", "created_time", "TIMESTAMP", "rec.get('inputTime')", "Mốc tạo/phân phối đơn"],
        ["5", "pickup_station", "TEXT", "rec.get('pickNetworkName')", "Tên bưu cục gốc lấy hàng"],
        ["6", "dispatch_code", "TEXT", "extract_ma10(terminalDispatchCode)", "Mã sortcode 10 ký tự tra cứu vùng"],
        ["7", "orders_num", "INT", "rec.get('packageNumber')", "Số kiện"],
        ["8", "orders_weight", "FLOAT", "rec.get('packageChargeWeight')", "Trọng lượng thô (gram)"],
        ["9", "pickup_station2", "TEXT", "rec.get('realPickNetworkName')", "Bưu cục lấy thực tế"],
        ["10", "pickup_time", "TIMESTAMP", "rec.get('pickTime')", "Thời gian bưu cục quét pickup"],
        ["11", "pickup_ontime", "TEXT", "Đánh giá đúng giờ pickup", "Đánh giá KPI Pickup"],
        ["12", "areacode", "TEXT", "rec.get('proxyAreaCode')", "Mã area phân phối"],
        ["13", "flowtypedesc", "TEXT", "rec.get('flowTypeDesc')", "Loại luồng hàng"],
        ["14", "next_station", "TEXT", "Tra cứu valid.csv theo dispatch_code", "Tên bưu cục đích hiển thị chuẩn"],
        ["15", "round", "TEXT", "Tra cứu valid.csv theo dispatch_code", "Tuyến giao hàng"],
        ["16", "rank", "TEXT", "Tra cứu valid.csv theo dispatch_code", "Phân cấp bưu cục (Rank 1, 2, 3)"],
        ["17", "inbound_scandate", "TIMESTAMP", "Khớp từ ib_scan_map[tracking]", "Mốc thời gian nhập kho HUB"],
        ["18", "outbound_scandate", "TIMESTAMP", "Khớp từ ob_map[tracking]", "Mốc thời gian xuất kho HUB"],
        ["19", "arrival_scandate", "TIMESTAMP", "Khớp từ arr_scan_map[tracking]", "Mốc thời gian đến bưu cục gửi"],
        ["20", "trip_code", "TEXT", "Khớp từ ib_trip_map / arr_trip_map", "Mã chuyến xe vận chuyển"],
        ["21", "transporing_time", "TIMESTAMP", "Tra từ ttm[trip_code]['transporing_time']", "Mốc giờ xe bắt đầu chạy"],
        ["22", "transported_time", "TIMESTAMP", "Tra từ ttm[trip_code]['transported_time']", "Mốc giờ xe cập bến"],
        ["23", "dispatch_actual", "TEXT", "Mã phân loại thực tế", "Đơn vị thực tế nhận"],
        ["24", "operation_date_created", "DATE", "get_op_date(created_time)", "Ngày vận hành tạo đơn (ca 06h)"],
        ["25", "operation_date_inbound", "DATE", "get_op_date(inbound_scandate)", "Ngày vận hành nhập kho HUB"],
        ["26", "is_backlog", "INT", "1 nếu (has_inbound AND NOT has_outbound)", "Cờ báo đơn đang tồn kho HUB"],
        ["27", "is_active", "INT", "0 nếu đã Outbound, ngược lại 1", "Cờ đơn đang hoạt động (chưa rời kho)"],
        ["28", "is_transit", "INT", "1 nếu (has_in AND NOT has_out AND has_arr)", "Cờ đơn đang trên xe trung chuyển"]
    ]
    tbl_pg = doc.add_table(rows=1, cols=5)
    format_table_headers_and_cells(tbl_pg, pg_widths, pg_headers, pg_data)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 5
    add_heading_1(doc, "5. NỘI DUNG CÁC FILE JSON & ÁNH XẠ LÊN REACT DASHBOARD")
    add_body_p(doc, "Chi tiết 11 file JSON chính thức được tạo ra từ PostgreSQL và cách Frontend React tiêu thụ dữ liệu:")

    json_headers = ["Tên File JSON", "Nội Dung Dữ Liệu Payload", "Component Frontend Tiêu Thụ", "Cách Hiển Thị / Ánh Xạ UI"]
    json_widths = [1.5, 2.2, 1.8, 1.7]
    json_data = [
        ["inbound.json / latest.json.gz", "Gom nhóm theo 15-tuple key + field is_north/region. Chứa volume, weight_ton, trip_code, drop_type, is_rebound...", "InboundDashboard.tsx (App.tsx giải nén gzip)", "Vẽ biểu đồ xu hướng theo giờ (Hourly Trend), thẻ KPI Forecast/Inbound/Weight, và Bảng Bưu cục gửi."],
        ["inventory.json", "Tồn kho theo từng ô Chute/Rack. Chứa zone, area_id, station_name, volume, weight_ton, capacity...", "App.tsx (Sơ đồ kho Master)", "Tô màu trực tiếp lên các ô Chute/Rack theo % sử dụng, hiển thị thông số ở Bảng Control Center."],
        ["outbound.json", "Sản lượng xuất kho HUB trong 2 ngày. Chứa zone, area_id, station_name, volume, weight_ton...", "App.tsx (Data Availability)", "Dữ liệu sẵn sàng trên GitHub Pages cho Outbound tracking."],
        ["backlog.json", "Đơn hàng đã Inbound nhưng chưa Outbound (tồn đọng). Chứa zone, area_id, station_name, volume...", "App.tsx (KPI Backlog)", "Cảnh báo màu đỏ ở các ô Chute có lượng đơn vượt sức chứa định mức."],
        ["arrival.json", "Tiến độ xe hàng đến kho theo khung giờ. Chứa scan_hour, total_orders, at_hub, not_hub...", "Arrival Monitor UI", "Đếm số đơn đã lên sàn HUB (at_hub) và số đơn còn nằm trên xe ngoài bãi (not_hub)."],
        ["heatmap.json", "Dict 24 slot giờ (00:00 - 23:00) đếm số lượt nhập kho trong ngày.", "HeatmapDashboard.tsx", "Vẽ biểu đồ nhiệt khung giờ cao điểm nhập kho trong ca làm việc."],
        ["linehaul.json & truck_eta.json", "Danh sách chuyến xe tải đang di chuyển. Chứa plate_number, trip_code, orders_count, weight_ton, eta...", "InboundDashboard.tsx (ETA Card)", "Thẻ KPI 'Inbound Truck ETA - HCM HUB' đếm số xe đang chạy đến HUB và tổng số đơn/tấn sắp đổ về sàn."],
        ["last_update.json", "Metadata thời gian đồng bộ: last_update, active_date, total_records, sync_success.", "Header Bar (Toàn bộ UI)", "Hiển thị mốc thời gian cập nhật dữ liệu gần nhất ở góc trên màn hình."]
    ]
    tbl_json = doc.add_table(rows=1, cols=4)
    format_table_headers_and_cells(tbl_json, json_widths, json_headers, json_data)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 6
    add_heading_1(doc, "6. LUỒNG ÁNH XẠ KỸ THUẬT LÊN FRONTEND REACT")
    add_body_p(doc, "Quy trình xử lý dữ liệu khi người dùng truy cập Dashboard trên trình duyệt:")

    add_bullet_p(doc, "1. Phát hiện Môi trường (Dual Hostname Fetching):", "App.tsx kiểm tra window.location.hostname. Nếu trên 'github.io' sẽ đọc file từ Raw GitHub CDN kèm parameter timestamp (?t=Date.now()) để chống browser cache. Nếu local dev sẽ đọc file local ./data/.")
    add_bullet_p(doc, "2. Chuẩn hóa Ánh xạ Key (Dual Key Mapping):", "Hàm fetchInboundSheetData() tạo keyMap nhân bản dữ liệu: 'station_name' ↔ 'Bưu cục', 'status' ↔ 'Trạng thái', 'volume' ↔ 'Volume', 'weight_ton' ↔ 'Weight', 'inbound_hour' ↔ 'Inbound Hour'. Giúp vừa hỗ trợ code React chuẩn vừa tương thích UI legacy.")
    add_bullet_p(doc, "3. Chuẩn hóa Enum Trạng thái & Drop Type:", "Hàm normalizeStatus() đổi tất cả alias ('at_hub', 'Đang trên bãi'...) về 'Inbound'. Hàm normalizeDropType('') giữ nguyên chuỗi rỗng nếu rỗng, tránh gán nhầm đơn rớt.")
    add_bullet_p(doc, "4. Render Giao diện & Biểu đồ:", "Dữ liệu được truyền vào InboundDashboard.tsx để tính toán tổng đơn, tổng tấn, vẽ 4 đường xu hướng theo giờ (Chart.js / SVG) và dựng bảng Bưu cục nộp hàng.")

    # Section 7
    add_heading_1(doc, "7. CƠ CHẾ TỰ ĐỘNG HÓA TIẾN TRÌNH & GIÁM SÁT NHẬT KÝ (AUTO-SYNC & LOGGING)")
    add_body_p(doc, "Hệ thống triển khai cơ chế đồng bộ tự động 24/7 và giám sát nhật ký vận hành đa tầng như sau:")

    sync_headers = ["Thành Phần Vận Hành", "Loại / Công Cụ", "Đường Dẫn / Tên Cấu Hình", "Mô Tả & Logic Kỹ Thuật"]
    sync_widths = [1.8, 1.3, 2.2, 1.7]
    sync_data = [
        ["Scheduled Task", "Windows Task Scheduler", "Task Name: Sync_Postgre_30m", "Tự động kích hoạt định kỳ mỗi 30 phút (RepetitionInterval 30m, vô thời hạn). Chạy dưới quyền tài khoản người dùng đăng nhập."],
        ["Batch Launcher", "Windows Batch Script", "backend_sync/run_sync_postgre.bat", "Thiết lập PYTHONIOENCODING=utf-8, chuyển Cwd về thư mục backend_sync, gọi py sync_postgre.py và append output vào log."],
        ["PowerShell Setup", "PowerShell Script", "backend_sync/setup_sync_postgre_task.ps1", "Kịch bản tự động hóa việc đăng ký Task vào Windows Task Scheduler bằng cmdlet Register-ScheduledTask."],
        ["Execution Log File", "Flat File Log", "backend_sync/sync_postgre.log", "Ghi vết chi tiết từng phiên đồng bộ: mốc thời gian bắt đầu/kết thúc, các lỗi ngoại lệ (exception) và kết quả nén JSON/Git Push."],
        ["System Event Log", "Windows Event Viewer", "taskschd.msc (History Tab)", "Theo dõi lịch sử kích hoạt của Windows, trạng thái Task (Ready/Running), mã thoát (Task Result code 0x0/Last Run Time)."]
    ]
    tbl_sync = doc.add_table(rows=1, cols=4)
    format_table_headers_and_cells(tbl_sync, sync_widths, sync_headers, sync_data)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Footer / Sign off
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_foot = p_footer.add_run("--- HẾT BỘ TÀI LIỆU KỸ THUẬT ---")
    r_foot.font.name = 'Segoe UI'
    r_foot.font.size = Pt(10)
    r_foot.font.bold = True
    r_foot.font.color.rgb = RGBColor(120, 120, 120)

    base_dir = os.path.dirname(os.path.abspath(__file__))
    v2_path = os.path.join(base_dir, "Tai_Lieu_Ky_Thuat_System_Architecture_JFS_HUB_v2.docx")
    v1_path = os.path.join(base_dir, "Tai_Lieu_Ky_Thuat_System_Architecture_JFS_HUB.docx")
    
    # Save v2 directly
    doc.save(v2_path)
    print(f"SUCCESS: Da tao thanh cong file Word v2 tai: {v2_path}")
    
    # Try saving v1 as well if accessible
    try:
        doc.save(v1_path)
        print(f"SUCCESS: Da tao thanh cong file Word v1 tai: {v1_path}")
    except Exception as e:
        print(f"INFO: File v1 khong ghi duoc (dang mo hoac dang lock): {e}")

    return v2_path

if __name__ == '__main__':
    generate_docx()
